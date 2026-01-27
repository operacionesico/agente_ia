"""
Script para inyectar datos y contenido IA en plantillas de Word.
Lee configuraciones desde Excel en '1. Data/', procesa plantillas en '2. Plantilla/',
y guarda resultados en '3. Inyectado/'.
"""

import os
import glob
from docx import Document
import openpyxl
import re
from datetime import datetime
from gemini_client import GeminiClient
from config_auditoria import generar_contexto_base
from lector_informacion import leer_informacion_empresa

# Configuración de rutas (AJUSTADO A TUS NOMBRES DE CARPETAS)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, '1. Data')
PLANTILLA_DIR = os.path.join(BASE_DIR, '2. Plantilla')
INYECTADO_DIR = os.path.join(BASE_DIR, '3. Inyectado')
IMAGENES_DIR = os.path.join(BASE_DIR, '5. IMAGENES')

def leer_datos_excel():
    """
    Busca el primer archivo Excel en '1. Data/' y lee las configuraciones.
    
    Nueva estructura simplificada (3 columnas):
    - Columna A: Campo (nombre de la variable)
    - Columna B: Valor (dato estático O prompt si es IA)
    - Columna C: Campo Generado (etiqueta)
    
    Detección automática:
    - Si Campo Generado empieza con {{IA: → Valor contiene el prompt
    - Si no → Valor contiene el dato estático
    
    Retorna dos diccionarios: datos_estaticos y datos_ia
    """
    # Buscar archivos .xlsx
    archivos_excel = glob.glob(os.path.join(DATA_DIR, '*.xlsx'))
    
    if not archivos_excel:
        print(f"❌ No se encontró ningún archivo Excel en la carpeta '{DATA_DIR}'")
        return None, None
    
    archivo_excel = archivos_excel[0]
    print(f"📄 Leyendo datos desde: {os.path.basename(archivo_excel)}")
    
    wb = openpyxl.load_workbook(archivo_excel, data_only=True)
    ws = wb.active
    
    datos_estaticos = {}
    datos_ia = {}
    
    print("   Leyendo filas...")
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not row[0]: continue  # Saltar filas vacías
        
        campo = str(row[0]).strip()
        valor = row[1] if row[1] is not None else ""
        campo_generado = str(row[2]).strip() if len(row) > 2 and row[2] else ""
        
        # Detección automática según el patrón del Campo Generado
        if campo_generado.startswith("{{IA:"):
            # Es una etiqueta IA, el valor contiene el prompt
            datos_ia[campo] = valor
            print(f"   - Fila {i}: {campo} [IA] → Prompt configurado")
        else:
            # Es una etiqueta estática, el valor contiene el dato
            datos_estaticos[campo] = valor
            print(f"   - Fila {i}: {campo} [STATIC] → {str(valor)[:50]}...")
    
    
    # Detectar normas automáticamente (campos que contengan "NORMA")
    normas = []
    for campo, valor in datos_estaticos.items():
        if 'NORMA' in campo.upper() and valor:
            normas.append(str(valor))
    
    wb.close()
    return datos_estaticos, datos_ia, normas

def procesar_parrafo(parrafo, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, memoria_respuestas):
    """
    Reemplaza etiquetas en un párrafo de Word.
    Soporta {{ETIQUETA}} (estáticas), {{IA:NOMBRE}} (IA) y {{IMG:CAMPO}} (imágenes)
    Incluye memoria de respuestas previas para coherencia.
    """
    from docx.shared import Inches
    texto = parrafo.text
    cambios = False
    
    # 1. Primero reemplazar etiquetas IA: {{IA:NOMBRE}}
    # Esto debe ir primero para no confundirlas con las estáticas
    etiquetas_ia = re.findall(r'\{\{IA:([^}]+)\}\}', texto)
    
    for nombre_prompt in etiquetas_ia:
        if nombre_prompt in datos_ia:
            print(f"   🤖 Generando contenido IA para: {nombre_prompt}...")
            prompt = datos_ia[nombre_prompt]
            
            # Generar respuesta sin memoria acumulativa (datos del Excel y RAG son suficientes)
            respuesta = cliente_gemini.generar_texto(prompt, datos_estaticos, contexto_sistema)
            
            if respuesta.startswith("[ERROR"):
                print(f"      ⚠️ {respuesta}")
            else:
                # Guardar respuesta en memoria
                memoria_respuestas.append(f"[{nombre_prompt}]\n{respuesta}")
            
            marcador = f"{{{{IA:{nombre_prompt}}}}}"
            if marcador in texto:
                texto = texto.replace(marcador, respuesta)
                cambios = True
    
    # 2. Procesar etiquetas de imagen: {{IMG:CAMPO}} - DESACTIVADO TEMPORALMENTE
    # etiquetas_img = re.findall(r'\{\{IMG:([^}]+)\}\}', texto)
    # (código comentado - funcionalidad de imágenes pendiente)
    # for campo_img in etiquetas_img:
    #     if campo_img in datos_estaticos:
    #         nombre_archivo = str(datos_estaticos[campo_img])
    #         ruta_imagen = os.path.join(IMAGENES_DIR, nombre_archivo)
            
    #         if os.path.exists(ruta_imagen):
    #             # Eliminar el marcador de imagen del texto
    #             marcador_img = f"{{{{IMG:{campo_img}}}}}"
    #             texto = texto.replace(marcador_img, "")
    #             cambios = True
                
    #             # Limpiar el párrafo y agregar la imagen
    #             parrafo.clear()
    #             run = parrafo.add_run()
                
    #             try:
    #                 run.add_picture(ruta_imagen, width=Inches(2.0))
    #                 print(f"   🖼️  Imagen insertada: {nombre_archivo}")
    #             except Exception as e:
    #                 parrafo.text = f"[Error insertando imagen: {e}]"
    #                 print(f"   ⚠️  Error insertando imagen {nombre_archivo}: {e}")
    #         else:
    #             texto = texto.replace(f"{{{{IMG:{campo_img}}}}}", f"[Imagen no encontrada: {nombre_archivo}]")
    #             print(f"   ⚠️  Imagen no encontrada: {ruta_imagen}")
    #             cambios = True
    
    # 3. Luego reemplazar etiquetas estáticas: {{CAMPO}}
    # Buscar todas las etiquetas {{ALGO}} que NO sean {{IA:...}}
    etiquetas_estaticas = re.findall(r'\{\{(?!IA:)([^}]+)\}\}', texto)
    
    for etiqueta in etiquetas_estaticas:
        if etiqueta in datos_estaticos:
            valor = str(datos_estaticos[etiqueta])
            marcador = f"{{{{{etiqueta}}}}}"
            if marcador in texto:
                texto = texto.replace(marcador, valor)
                cambios = True
    
    if cambios:
        # 4. LIMPIEZA POST-PROCESAMIENTO: Eliminar comas y espacios sobrantes
        # Esto soluciona casos como: "ISO 9001, ISO 45001,  ,  " -> "ISO 9001, ISO 45001"
        
        # Limpiar comas múltiples con o sin espacios: ",  ," -> ","
        texto = re.sub(r',\s*,+', ',', texto)
        
        # Limpiar comas al final (antes de punto, salto de línea, etc)
        texto = re.sub(r',\s*([.\n)])', r'\1', texto)
        
        # Limpiar espacios múltiples
        texto = re.sub(r'\s{2,}', ' ', texto)
        
        # Limpiar coma+espacio al final de línea
        texto = re.sub(r',\s*$', '', texto, flags=re.MULTILINE)
        
        parrafo.text = texto

def procesar_word():
    print("="*60)
    print("INYECTOR DE WORD CON IA - INICIANDO v1.0")
    print("="*60)

    # 1. Inicializar cliente Gemini
    try:
        cliente_gemini = GeminiClient()
        print("✅ Cliente Gemini correecto")
    except Exception as e:
        print(f"❌ Error al iniciar Gemini: {e}")
        return

    # 2. Leer datos
    datos_estaticos, datos_ia, normas = leer_datos_excel()
    if not datos_estaticos:
        print("❌ No se pudieron cargar datos. Verifica la carpeta '1. Data'")
        return
    
    # 2.5 Leer información adicional de la empresa (RAG)
    contexto_empresa = leer_informacion_empresa()
    
    # 2.6 Generar contexto global de auditoría
    contexto_sistema = generar_contexto_base(normas, datos_estaticos, None, contexto_empresa)
    
    # 2.7 GUARDAR CONTEXTO EN ARCHIVO TXT PARA REVISIÓN
    ruta_contexto = os.path.join(INYECTADO_DIR, 'CONTEXTO_IA.txt')
    try:
        with open(ruta_contexto, 'w', encoding='utf-8') as f:
            f.write("="*80 + "\n")
            f.write("CONTEXTO COMPLETO ENVIADO A LA IA\n")
            f.write(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
            f.write("="*80 + "\n\n")
            f.write(contexto_sistema)
        print(f"💾 Contexto guardado en: {os.path.basename(ruta_contexto)}")
    except Exception as e:
        print(f"⚠️  No se pudo guardar contexto: {e}")
    
    if normas:
        print(f"📋 Normas detectadas para auditoría: {', '.join(normas)}")
    else:
        print("⚠️  No se detectaron normas en el Excel (campos NORMA1, NORMA2, etc.)")
        
    print(f"📊 Datos cargados: {len(datos_estaticos)} variables estáticas, {len(datos_ia)} prompts de IA")
    
    # 3. Buscar plantillas (excluir archivos temporales de Word)
    todas_plantillas = glob.glob(os.path.join(PLANTILLA_DIR, '*.docx'))
    plantillas = [p for p in todas_plantillas if not os.path.basename(p).startswith('~$')]
    
    if not plantillas:
        print(f"❌ No se encontraron plantillas .docx en '{PLANTILLA_DIR}'")
        return
        
    print(f"📂 Se encontraron {len(plantillas)} plantillas.")
    
    # 4. Procesar cada plantilla
    for ruta_plantilla in plantillas:
        nombre_archivo = os.path.basename(ruta_plantilla)
        print(f"\n🔄 Procesando plantilla: {nombre_archivo}")
        
        # Inicializar memoria de respuestas para esta plantilla
        memoria_respuestas = []
        
        try:
            doc = Document(ruta_plantilla)
            
            # Procesar párrafos del cuerpo
            for parrafo in doc.paragraphs:
                procesar_parrafo(parrafo, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, memoria_respuestas)
                
            # Procesar tablas (celda por celda)
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for parrafo in celda.paragraphs:
                            procesar_parrafo(parrafo, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, memoria_respuestas)
            
            # 5. Guardar resultado
            if not os.path.exists(INYECTADO_DIR):
                os.makedirs(INYECTADO_DIR)
            
            # Generar nombre de salida (quitar _ inicial si existe)
            nombre_base = os.path.splitext(nombre_archivo)[0]  # Sin extensión
            if nombre_base.startswith('_'):
                nombre_base = nombre_base[1:]  # Quitar el primer carácter (_)
            
            nombre_salida = f"{nombre_base}.docx"
            ruta_salida = os.path.join(INYECTADO_DIR, nombre_salida)
            doc.save(ruta_salida)
            print(f"✅ Documento guardado exitosamente en: 3. Inyectado/{nombre_salida}")
            
            # Guardar memoria completa para revisión
            if memoria_respuestas:
                memoria_file = os.path.join(INYECTADO_DIR, f'MEMORIA_{nombre_salida.replace(".docx", ".txt")}')
                try:
                    with open(memoria_file, 'w', encoding='utf-8') as f:
                        f.write("MEMORIA COMPLETA DE GENERACIONES IA\n")
                        f.write("="*80 + "\n\n")
                        for i, resp in enumerate(memoria_respuestas, 1):
                            f.write(f"\n{i}. {resp}\n")
                            f.write("-"*80 + "\n")
                    print(f"💾 Memoria guardada en: MEMORIA_{nombre_salida.replace('.docx', '.txt')}")
                except Exception as e:
                    print(f"⚠️  No se pudo guardar memoria: {e}")
            
        except Exception as e:
            print(f"❌ Error procesando {nombre_archivo}: {str(e)}")

    print("\n" + "="*60)
    print("PROCESO COMPLETADO")
    print("="*60)

if __name__ == "__main__":
    procesar_word()