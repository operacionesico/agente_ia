"""
Procesador adaptado para Streamlit - Trabaja con archivos en memoria
Soporta Word (.docx) y Excel (.xlsx)
"""

import io
import tempfile
import os
import shutil
from docx import Document
import openpyxl
from datetime import datetime
import re

from gemini_client import GeminiClient
from config_auditoria import generar_contexto_base

def procesar_documentos_streamlit(data_file, plantillas, docs_empresa=None, progress_callback=None):
    """
    Procesa documentos usando archivos en memoria (UploadedFile de Streamlit).
    Soporta plantillas Word (.docx) y Excel (.xlsx).
    
    Args:
        data_file: UploadedFile con DATA.xlsx
        plantillas: Lista de UploadedFile con plantillas .docx o .xlsx
        docs_empresa: Lista de UploadedFile con documentos empresa (opcional)
        progress_callback: Función callback(progress, mensaje) para actualizar UI
    
    Returns:
        Lista de diccionarios {'nombre': str, 'contenido': bytes}
    """
    
    def update_progress(percent, message):
        if progress_callback:
            progress_callback(min(percent, 100), message)
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        update_progress(15, "Leyendo datos del Excel...")
        
        datos_estaticos, datos_ia, normas = leer_datos_excel_memoria(data_file)
        
        update_progress(25, "Inicializando cliente IA...")
        
        cliente_gemini = GeminiClient()
        
        update_progress(30, "Procesando información empresarial...")
        
        contexto_empresa = ""
        if docs_empresa:
            contexto_empresa = procesar_docs_empresa_memoria(docs_empresa, temp_dir)
        
        update_progress(35, "Generando contexto base...")
        
        contexto_sistema = generar_contexto_base(normas, datos_estaticos, None, contexto_empresa)
        
        documentos_generados = []
        num_plantillas = len(plantillas)
        
        memoria_total = []
        tiene_word = any(p.name.endswith('.docx') for p in plantillas)
        tiene_excel = any(p.name.endswith('.xlsx') for p in plantillas)
        
        for idx, plantilla_file in enumerate(plantillas):
            progress_inicio = 40 + (idx * 50 // num_plantillas)
            progress_fin = 40 + ((idx + 1) * 50 // num_plantillas)
            
            update_progress(progress_inicio, f"Procesando plantilla {idx+1}/{num_plantillas}: {plantilla_file.name}...")
            
            es_excel = plantilla_file.name.endswith('.xlsx')
            
            if es_excel:
                doc_generado, respuestas = procesar_plantilla_excel_memoria(
                    plantilla_file,
                    datos_estaticos,
                    datos_ia,
                    cliente_gemini,
                    contexto_sistema,
                    lambda p: update_progress(progress_inicio + int((progress_fin - progress_inicio) * p / 100), 
                                             f"Procesando {plantilla_file.name}... {p}%")
                )
                memoria_total.extend(respuestas)
            else:
                doc_generado, respuestas = procesar_plantilla_word_memoria(
                    plantilla_file,
                    datos_estaticos,
                    datos_ia,
                    cliente_gemini,
                    contexto_sistema,
                    lambda p: update_progress(progress_inicio + int((progress_fin - progress_inicio) * p / 100), 
                                             f"Procesando {plantilla_file.name}... {p}%")
                )
                memoria_total.extend(respuestas)
            
            nombre_salida = plantilla_file.name.replace('_', '').strip()
            
            output_buffer = io.BytesIO()
            doc_generado.save(output_buffer)
            output_buffer.seek(0)
            
            documentos_generados.append({
                'nombre': nombre_salida,
                'contenido': output_buffer.getvalue(),
                'tipo': 'documento'
            })
            
            update_progress(progress_fin, f"Completado: {nombre_salida}")
        
        update_progress(92, "Generando archivos de contexto...")
        
        if tiene_word:
            documentos_generados.append({
                'nombre': 'CONTEXTO_IA.txt',
                'contenido': contexto_sistema.encode('utf-8'),
                'tipo': 'contexto'
            })
        
        if tiene_excel:
            documentos_generados.append({
                'nombre': 'CONTEXTO_IA_EXCEL.txt',
                'contenido': contexto_sistema.encode('utf-8'),
                'tipo': 'contexto'
            })
        
        update_progress(95, "Generando archivo de memoria...")
        
        if memoria_total:
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            memoria_txt = "\n\n".join(memoria_total)
            
            if tiene_word and tiene_excel:
                nombre_memoria = f"MEMORIA_WORD_EXCEL_{timestamp}.txt"
            elif tiene_word:
                nombre_memoria = f"MEMORIA_WORD_{timestamp}.txt"
            else:
                nombre_memoria = f"MEMORIA_EXCEL_{timestamp}.txt"
            
            documentos_generados.append({
                'nombre': nombre_memoria,
                'contenido': memoria_txt.encode('utf-8'),
                'tipo': 'memoria'
            })
        
        update_progress(98, "Finalizando...")
        
        return documentos_generados
        
    finally:
        try:
            shutil.rmtree(temp_dir)
        except:
            pass

def leer_datos_excel_memoria(excel_file):
    """Lee datos del Excel desde un UploadedFile."""
    
    wb = openpyxl.load_workbook(io.BytesIO(excel_file.read()), data_only=True)
    ws = wb.active
    
    datos_estaticos = {}
    datos_ia = {}
    
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row[0]:
            continue
        
        campo = str(row[0]).strip()
        valor = row[1] if row[1] is not None else ""
        campo_generado = str(row[2]).strip() if len(row) > 2 and row[2] else ""
        
        if campo_generado.startswith("{{IA:"):
            datos_ia[campo] = valor
        else:
            datos_estaticos[campo] = valor
    
    normas = []
    for campo, valor in datos_estaticos.items():
        if 'NORMA' in campo.upper() and valor and not campo.endswith('_RESUMEN'):
            normas.append(str(valor))
    
    wb.close()
    
    return datos_estaticos, datos_ia, normas

def procesar_docs_empresa_memoria(docs_files, temp_dir):
    """Procesa documentos de empresa desde UploadedFiles."""
    from PyPDF2 import PdfReader
    import openpyxl
    
    contexto = "\n\n═══════════════════════════════════════════════════════════════════════\n"
    contexto += "INFORMACIÓN ADICIONAL DE LA EMPRESA:\n"
    contexto += "═══════════════════════════════════════════════════════════════════════\n\n"
    
    for doc_file in docs_files:
        temp_path = os.path.join(temp_dir, doc_file.name)
        with open(temp_path, 'wb') as f:
            f.write(doc_file.read())
        
        try:
            if doc_file.name.endswith('.txt'):
                with open(temp_path, 'r', encoding='utf-8') as f:
                    contenido = f.read()
            elif doc_file.name.endswith('.docx'):
                doc = Document(temp_path)
                contenido = '\n'.join([p.text for p in doc.paragraphs])
            elif doc_file.name.endswith('.pdf'):
                reader = PdfReader(temp_path)
                contenido = ""
                for pagina in reader.pages:
                    contenido += pagina.extract_text() + "\n"
            elif doc_file.name.endswith('.xlsx'):
                wb = openpyxl.load_workbook(temp_path, data_only=True)
                contenido = ""
                for hoja in wb.sheetnames:
                    ws = wb[hoja]
                    contenido += f"\n--- Hoja: {hoja} ---\n"
                    for fila in ws.iter_rows(values_only=True):
                        valores = [str(v) for v in fila if v is not None]
                        if valores:
                            contenido += " | ".join(valores) + "\n"
                wb.close()
            else:
                contenido = f"[Documento: {doc_file.name}]"
            
            contexto += f"--- {doc_file.name} ---\n{contenido[:2000]}\n\n"
        except Exception as e:
            contexto += f"--- {doc_file.name} ---\n[Error procesando: {str(e)}]\n\n"
    
    return contexto


def procesar_plantilla_word_memoria(plantilla_file, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, progress_callback=None):
    """Procesa una plantilla Word desde UploadedFile."""
    
    doc = Document(io.BytesIO(plantilla_file.read()))
    
    memoria_respuestas = []
    
    total_parrafos = len(doc.paragraphs)
    
    for idx, parrafo in enumerate(doc.paragraphs):
        if progress_callback:
            progress_callback(int((idx / total_parrafos) * 100))
        
        procesar_parrafo_streamlit(parrafo, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, memoria_respuestas)
    
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    procesar_parrafo_streamlit(parrafo, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, memoria_respuestas)
    
    return doc, memoria_respuestas

def procesar_plantilla_excel_memoria(plantilla_file, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, progress_callback=None):
    """Procesa una plantilla Excel desde UploadedFile."""
    
    wb = openpyxl.load_workbook(io.BytesIO(plantilla_file.read()))
    
    memoria_respuestas = []
    
    total_celdas = sum(len(list(ws.iter_rows())) for ws in wb.worksheets)
    celdas_procesadas = 0
    
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for celda in row:
                procesar_celda_streamlit(celda, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, memoria_respuestas)
                
                celdas_procesadas += 1
                if progress_callback and total_celdas > 0:
                    progress_callback(int((celdas_procesadas / total_celdas) * 100))
    
    return wb, memoria_respuestas


def procesar_parrafo_streamlit(parrafo, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, memoria_respuestas):
    """Procesa un párrafo individual de Word."""
    
    texto = parrafo.text
    cambios = False
    
    for campo, valor in datos_estaticos.items():
        etiqueta = f"{{{{{campo}}}}}"
        if etiqueta in texto:
            texto = texto.replace(etiqueta, str(valor))
            cambios = True
    
    etiquetas_ia = re.findall(r'\{\{IA:([^}]+)\}\}', texto)
    
    for nombre_prompt in etiquetas_ia:
        if nombre_prompt in datos_ia:
            prompt = datos_ia[nombre_prompt]
            
            contexto_sistema = contexto_sistema
            if memoria_respuestas:
                contexto_sistema += "\n\n" + "="*80 + "\n"
                contexto_sistema += "RESPUESTAS GENERADAS PREVIAMENTE (mantén coherencia con estos datos):\n"
                contexto_sistema += "="*80 + "\n\n"
                contexto_sistema += "\n\n".join(memoria_respuestas[-15:])
            
            respuesta = cliente_gemini.generar_texto(prompt, datos_estaticos, contexto_sistema)
            
            memoria_respuestas.append(f"[{nombre_prompt}]\n{respuesta}")
            
            texto = texto.replace(f"{{{{IA:{nombre_prompt}}}}}", respuesta)
            cambios = True
    
    if cambios:
        texto = re.sub(r',\s*,', ',', texto)
        texto = re.sub(r'\s+,', ',', texto)
        texto = re.sub(r',\s*\.', '.', texto)
        texto = re.sub(r'\s{2,}', ' ', texto)
        
        parrafo.text = texto.strip()

def procesar_celda_streamlit(celda, datos_estaticos, datos_ia, cliente_gemini, contexto_sistema, memoria_respuestas):
    """Procesa una celda individual de Excel."""
    
    if celda.value is None or not isinstance(celda.value, str):
        return
    
    texto = str(celda.value)
    cambios = False
    
    for campo, valor in datos_estaticos.items():
        etiqueta = f"{{{{{campo}}}}}"
        if etiqueta in texto:
            texto = texto.replace(etiqueta, str(valor))
            cambios = True
    
    etiquetas_ia = re.findall(r'\{\{IA:([^}]+)\}\}', texto)
    
    for nombre_prompt in etiquetas_ia:
        if nombre_prompt in datos_ia:
            prompt = datos_ia[nombre_prompt]
            
            contexto_sistema = contexto_sistema
            if memoria_respuestas:
                contexto_sistema += "\n\n" + "="*80 + "\n"
                contexto_sistema += "RESPUESTAS GENERADAS PREVIAMENTE (mantén coherencia con estos datos):\n"
                contexto_sistema += "="*80 + "\n\n"
                contexto_sistema += "\n\n".join(memoria_respuestas[-15:])
            
            respuesta = cliente_gemini.generar_texto(prompt, datos_estaticos, contexto_sistema)
            
            memoria_respuestas.append(f"[{nombre_prompt}]\n{respuesta}")
            
            texto = texto.replace(f"{{{{IA:{nombre_prompt}}}}}", respuesta)
            cambios = True
    
    if cambios:
        texto = re.sub(r',\s*,', ',', texto)
        texto = re.sub(r'\s+,', ',', texto)
        texto = re.sub(r',\s*\.', '.', texto)
        texto = re.sub(r'\s{2,}', ' ', texto)
        
        celda.value = texto.strip()
